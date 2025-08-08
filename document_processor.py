import os
import json
import re
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
from datetime import datetime

# Try to import optional dependencies
try:
    import PyPDF2
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False
    
try:
    import docx
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

class AdvancedDocumentProcessor:
    """Enhanced document processor with better structure preservation"""
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.html']

    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text with metadata and structure preservation"""
        if not HAS_PDF_SUPPORT:
            return {'text': '', 'pages': [], 'metadata': {}, 'structure_type': 'pdf', 'error': 'PDF support not available'}
            
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                metadata = {
                    'title': pdf_reader.metadata.title if pdf_reader.metadata else None,
                    'author': pdf_reader.metadata.author if pdf_reader.metadata else None,
                    'total_pages': len(pdf_reader.pages),
                    'creation_date': str(pdf_reader.metadata.creation_date) if (pdf_reader.metadata and pdf_reader.metadata.creation_date) else None
                }

                # Extract text with page information
                pages_text = []
                full_text = ""

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    pages_text.append({
                        'page_number': page_num + 1,
                        'text': page_text,
                        'char_count': len(page_text)
                    })
                    full_text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"

                return {
                    'text': full_text,
                    'pages': pages_text,
                    'metadata': metadata,
                    'structure_type': 'pdf'
                }
        except Exception as e:
            print(f"Error extracting PDF {pdf_path}: {str(e)}")
            return {'text': '', 'pages': [], 'metadata': {}, 'structure_type': 'pdf'}

    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract text with document structure preservation"""
        if not HAS_DOCX_SUPPORT:
            return {'text': '', 'paragraphs': [], 'tables': [], 'metadata': {}, 'structure_type': 'docx', 'error': 'DOCX support not available'}
            
        try:
            doc = docx.Document(docx_path)

            # Extract paragraphs with styles
            structured_content = []
            full_text = ""

            for para in doc.paragraphs:
                para_data = {
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'level': self._get_heading_level(para.style.name) if para.style else 0
                }
                structured_content.append(para_data)

                # Add structure markers for headings
                if para_data['level'] > 0:
                    full_text += f"\n{'#' * para_data['level']} {para.text}\n"
                else:
                    full_text += f"{para.text}\n"

            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text += row_text + "\n"
                tables_data.append(table_text)
                full_text += f"\n--- TABLE ---\n{table_text}\n"

            return {
                'text': full_text,
                'paragraphs': structured_content,
                'tables': tables_data,
                'metadata': {'total_paragraphs': len(structured_content), 'total_tables': len(tables_data)},
                'structure_type': 'docx'
            }
        except Exception as e:
            print(f"Error extracting DOCX {docx_path}: {str(e)}")
            return {'text': '', 'paragraphs': [], 'tables': [], 'metadata': {}, 'structure_type': 'docx'}

    def _get_heading_level(self, style_name: str) -> int:
        """Determine heading level from style name"""
        if not style_name:
            return 0

        heading_map = {
            'Title': 1,
            'Heading 1': 1,
            'Heading 2': 2,
            'Heading 3': 3,
            'Heading 4': 4,
            'Heading 5': 5,
            'Heading 6': 6
        }

        return heading_map.get(style_name, 0)

    def extract_text_from_txt(self, txt_path: str) -> Dict[str, Any]:
        """Extract text with basic structure detection"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()

            # Detect basic structure (markdown-like headers)
            lines = text.split('\n')
            structured_content = []

            for line in lines:
                if line.strip().startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    structured_content.append({
                        'text': line.strip('#').strip(),
                        'type': 'heading',
                        'level': min(level, 6)
                    })
                elif line.strip():
                    structured_content.append({
                        'text': line,
                        'type': 'paragraph',
                        'level': 0
                    })

            return {
                'text': text,
                'structure': structured_content,
                'metadata': {'total_lines': len(lines)},
                'structure_type': 'txt'
            }
        except Exception as e:
            print(f"Error extracting TXT {txt_path}: {str(e)}")
            return {'text': '', 'structure': [], 'metadata': {}, 'structure_type': 'txt'}

class SmartChunkManager:
    """Advanced chunking with context preservation"""
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def create_smart_chunks(self, document_data: Dict[str, Any], document_name: str) -> List[Dict[str, Any]]:
        """Create intelligent chunks based on document structure"""
        text = document_data['text']
        structure_type = document_data.get('structure_type', 'unknown')

        chunks = []

        if structure_type == 'pdf' and 'pages' in document_data:
            chunks = self._chunk_pdf_by_pages(document_data, document_name)
        elif structure_type == 'docx' and 'paragraphs' in document_data:
            chunks = self._chunk_docx_by_structure(document_data, document_name)
        else:
            chunks = self._chunk_generic_text(text, document_name)

        # Add semantic metadata to chunks with enhanced categorization
        enhanced_chunks = []
        for i, chunk in enumerate(chunks):
            enhanced_chunk = chunk.copy()
            enhanced_chunk.update({
                'semantic_type': self._classify_chunk_type(chunk['text']),
                'content_categories': self._extract_content_categories(chunk['text']),
                'has_numbers': bool(re.search(r'\d+', chunk['text'])),
                'has_currency': bool(re.search(r'[\$\€\£\₹]|\b(?:dollar|euro|pound|rupee)s?\b', chunk['text'], re.IGNORECASE)),
                'has_dates': bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\b', chunk['text'], re.IGNORECASE)),
                'has_percentages': bool(re.search(r'\d+\.?\d*\s*%', chunk['text'])),
                'has_medical_terms': self._has_medical_terms(chunk['text']),
                'has_legal_terms': self._has_legal_terms(chunk['text']),
                'word_count': len(chunk['text'].split()),
                'chunk_position': i / len(chunks),  # Relative position in document
                'urgency_indicators': self._extract_urgency_indicators(chunk['text']),
                'exclusion_indicators': self._extract_exclusion_indicators(chunk['text'])
            })
            enhanced_chunks.append(enhanced_chunk)

        return enhanced_chunks

    def _chunk_pdf_by_pages(self, document_data: Dict[str, Any], document_name: str) -> List[Dict[str, Any]]:
        """Chunk PDF content by pages with overlap"""
        chunks = []
        pages = document_data['pages']
        
        for page_data in pages:
            page_text = page_data['text']
            if len(page_text.strip()) > 0:
                chunks.append({
                    'text': page_text,
                    'source': document_name,
                    'page_number': page_data['page_number'],
                    'chunk_type': 'page',
                    'metadata': {
                        'page_number': page_data['page_number'],
                        'char_count': page_data['char_count']
                    }
                })
        
        return chunks

    def _chunk_docx_by_structure(self, document_data: Dict[str, Any], document_name: str) -> List[Dict[str, Any]]:
        """Chunk DOCX content by structure"""
        chunks = []
        current_chunk = ""
        current_metadata = {}
        
        for para in document_data['paragraphs']:
            if para['level'] > 0 and current_chunk:
                # Start new chunk at heading
                chunks.append({
                    'text': current_chunk.strip(),
                    'source': document_name,
                    'chunk_type': 'structured',
                    'metadata': current_metadata
                })
                current_chunk = para['text'] + "\n"
                current_metadata = {'heading_level': para['level']}
            else:
                current_chunk += para['text'] + "\n"
                
            # Split if chunk gets too large
            if len(current_chunk) > self.chunk_size:
                chunks.append({
                    'text': current_chunk.strip(),
                    'source': document_name,
                    'chunk_type': 'structured',
                    'metadata': current_metadata
                })
                current_chunk = ""
                current_metadata = {}
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'source': document_name,
                'chunk_type': 'structured',
                'metadata': current_metadata
            })
        
        return chunks

    def _chunk_generic_text(self, text: str, document_name: str) -> List[Dict[str, Any]]:
        """Generic text chunking with overlap"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                'text': chunk_text,
                'source': document_name,
                'chunk_type': 'generic',
                'metadata': {
                    'start_index': i,
                    'end_index': min(i + self.chunk_size, len(words))
                }
            })
        
        return chunks

    def _classify_chunk_type(self, text: str) -> str:
        """Classify chunk content type"""
        text_lower = text.lower()
        
        if any(term in text_lower for term in ['table', 'chart', 'figure', '|', 'row', 'column']):
            return 'tabular'
        elif any(term in text_lower for term in ['section', 'chapter', 'article', 'clause']):
            return 'structural'
        elif any(term in text_lower for term in ['policy', 'coverage', 'premium', 'deductible']):
            return 'policy_content'
        elif any(term in text_lower for term in ['terms', 'conditions', 'agreement', 'contract']):
            return 'legal_content'
        else:
            return 'general'

    def _extract_content_categories(self, text: str) -> List[str]:
        """Extract multiple content categories for better search"""
        categories = []
        text_lower = text.lower()

        category_patterns = {
            'eligibility_criteria': ['eligible', 'eligibility', 'qualify', 'qualification', 'criteria', 'requirements'],
            'waiting_periods': ['waiting period', 'wait', 'minimum duration', 'cooling period', 'pre-existing'],
            'coverage_limits': ['maximum', 'limit', 'cap', 'upto', 'not exceeding', 'ceiling'],
            'deductibles': ['deductible', 'co-payment', 'copay', 'self-payment', 'excess'],
            'geographical': ['location', 'network', 'hospital', 'city', 'state', 'region'],
            'age_related': ['age', 'years old', 'minor', 'senior', 'adult', 'child'],
            'emergency': ['emergency', 'urgent', 'critical', 'immediate', 'acute'],
            'preventive': ['preventive', 'preventative', 'routine', 'screening', 'checkup'],
            'chronic_conditions': ['chronic', 'diabetes', 'hypertension', 'cancer', 'heart disease'],
            'maternity': ['maternity', 'pregnancy', 'childbirth', 'delivery', 'prenatal'],
            'dental': ['dental', 'teeth', 'oral', 'orthodontic', 'periodontal'],
            'mental_health': ['mental health', 'psychological', 'psychiatric', 'counseling', 'therapy']
        }

        for category, keywords in category_patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                categories.append(category)

        return categories

    def _has_medical_terms(self, text: str) -> bool:
        """Check if text contains medical terminology"""
        medical_terms = [
            'surgery', 'procedure', 'treatment', 'diagnosis', 'medication', 'therapy',
            'hospital', 'clinic', 'doctor', 'physician', 'specialist', 'consultation',
            'disease', 'condition', 'symptom', 'patient', 'medical', 'clinical'
        ]
        text_lower = text.lower()
        return any(term in text_lower for term in medical_terms)

    def _has_legal_terms(self, text: str) -> bool:
        """Check if text contains legal terminology"""
        legal_terms = [
            'contract', 'agreement', 'clause', 'terms', 'conditions', 'liability',
            'indemnity', 'breach', 'compliance', 'regulation', 'statute', 'law'
        ]
        text_lower = text.lower()
        return any(term in text_lower for term in legal_terms)

    def _extract_urgency_indicators(self, text: str) -> List[str]:
        """Extract urgency indicators from text"""
        indicators = []
        text_lower = text.lower()
        
        urgency_patterns = {
            'immediate': ['immediate', 'urgent', 'emergency', 'critical', 'acute'],
            'time_sensitive': ['within', 'before', 'deadline', 'expires', 'due date'],
            'conditional': ['must', 'required', 'mandatory', 'shall', 'obligated']
        }
        
        for urgency_type, keywords in urgency_patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                indicators.append(urgency_type)
        
        return indicators

    def _extract_exclusion_indicators(self, text: str) -> List[str]:
        """Extract exclusion indicators from text"""
        indicators = []
        text_lower = text.lower()
        
        exclusion_patterns = {
            'not_covered': ['not covered', 'excluded', 'exception', 'does not cover'],
            'limitations': ['limited to', 'maximum', 'up to', 'subject to'],
            'restrictions': ['restricted', 'prohibited', 'forbidden', 'not applicable']
        }
        
        for exclusion_type, keywords in exclusion_patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                indicators.append(exclusion_type)
        
        return indicators
